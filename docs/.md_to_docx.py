"""
Convert docs/scheduling-algorithm-walkthrough.md to a properly structured Word document.

Handles: H1/H2/H3 headings, paragraphs with bold (**), inline code (`),
bullet lists, numbered lists, tables, and horizontal rules.
Mixed Hebrew/English text inside LTR paragraphs renders correctly via Word's
built-in bidirectional handling.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

SRC = Path(r"c:\Projects\GardenManager\docs\scheduling-algorithm-walkthrough.md")
DST = Path(r"c:\Projects\GardenManager\docs\scheduling-algorithm-walkthrough.docx")


# ---- inline parsing -------------------------------------------------------

INLINE_TOKEN = re.compile(
    r"(\*\*[^*]+\*\*"        # **bold**
    r"|`[^`]+`"               # `code`
    r"|\[([^\]]+)\]\([^)]+\))"  # [text](link) -> use the text
)


def add_runs(paragraph, text: str):
    """Tokenize an inline string and append styled runs to the paragraph."""
    pos = 0
    for m in INLINE_TOKEN.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif token.startswith("["):
            # markdown link: keep visible text, drop URL
            link_text = m.group(2)
            paragraph.add_run(link_text)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ---- block parsing --------------------------------------------------------

BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
NUMBER_RE = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|?[\s\-:|]+\|[\s\-:|]+\|?\s*$")


def is_table_separator(line: str) -> bool:
    return bool(TABLE_SEP_RE.match(line)) and "-" in line


def split_table_row(line: str):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def add_horizontal_rule(doc):
    """Insert a thin bottom border on an empty paragraph as a horizontal rule."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def render_table(doc, rows):
    """rows[0] is header, rest are body."""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            # Replace existing paragraph contents
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, cell_text)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
    # Set first column narrow (the HC code column is short)
    if len(rows[0]) == 2:
        for row in table.rows:
            row.cells[0].width = Inches(0.7)
            row.cells[1].width = Inches(5.8)


def convert():
    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()

    # Tighten default style a bit
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # blank line
        if not stripped:
            i += 1
            continue

        # horizontal rule
        if stripped == "---":
            add_horizontal_rule(doc)
            i += 1
            continue

        # heading
        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2)
            level = min(level, 4)  # python-docx supports Heading 1..9
            h = doc.add_heading(level=level)
            add_runs(h, heading_text)
            i += 1
            continue

        # table: a row followed by a separator row
        if TABLE_ROW_RE.match(line) and i + 1 < n and is_table_separator(lines[i + 1]):
            header = split_table_row(line)
            i += 2  # skip header + separator
            body = []
            while i < n and TABLE_ROW_RE.match(lines[i]):
                body.append(split_table_row(lines[i]))
                i += 1
            render_table(doc, [header] + body)
            continue

        # bullet list
        if BULLET_RE.match(line):
            while i < n and BULLET_RE.match(lines[i]):
                bm = BULLET_RE.match(lines[i])
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, bm.group(2))
                i += 1
            continue

        # numbered list
        if NUMBER_RE.match(line):
            while i < n and NUMBER_RE.match(lines[i]):
                nm = NUMBER_RE.match(lines[i])
                p = doc.add_paragraph(style="List Number")
                add_runs(p, nm.group(3))
                i += 1
            continue

        # paragraph: collapse continuation lines until a blank or block boundary
        para_lines = [line]
        i += 1
        while i < n:
            nxt = lines[i]
            nxt_stripped = nxt.strip()
            if not nxt_stripped:
                break
            if HEADING_RE.match(nxt) or BULLET_RE.match(nxt) or NUMBER_RE.match(nxt):
                break
            if nxt_stripped == "---":
                break
            if TABLE_ROW_RE.match(nxt):
                break
            para_lines.append(nxt)
            i += 1
        para_text = " ".join(l.strip() for l in para_lines)
        p = doc.add_paragraph()
        add_runs(p, para_text)

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DST))
    print(f"wrote {DST} ({DST.stat().st_size} bytes)")


if __name__ == "__main__":
    try:
        convert()
    except Exception as e:
        print(f"error: {e}", file=sys.stderr)
        raise
